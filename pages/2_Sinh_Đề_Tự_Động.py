# file: sinh_de_kntc.py
import re
import io
import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from PIL import Image, ImageFile
import matplotlib
import matplotlib.pyplot as plt
import traceback
import logging
import unicodedata
import html

# --- Cấu hình logging (hữu ích khi debug) ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Matplotlib backend cho môi trường server ---
matplotlib.use("Agg")
plt.rcParams['mathtext.fontset'] = 'cm'
ImageFile.LOAD_TRUNCATED_IMAGES = True

st.set_page_config(page_title="Sinh Đề KNTC Tự Động", page_icon="📝", layout="wide")
st.title("📝 Sinh Đề Tự Động – Theo Ma Trận Đặc Tả Tối Giản")

# --- API KEY ---
api_key = st.secrets.get("GOOGLE_API_KEY", "")
if not api_key:
    api_key = st.text_input("Nhập Google API Key:", type="password")

# --- DỮ LIỆU MẪU (đã sửa trùng khóa) ---
lop_options = [
    "Lớp 1", "Lớp 2", "Lớp 3", "Lớp 4", "Lớp 5",
    "Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"
]

# Chuẩn: mỗi lớp có danh sách chương riêng
chuong_options = {
    "Lớp 1": ["Chủ đề 1: Các số đến 10", "Chủ đề 2: Các số đến 20", "Chủ đề 3: Các số đến 100", "Chủ đề 4: Hình học và đo lường", "Chủ đề 5: Giải toán"],
    "Lớp 2": ["Chủ đề 1: Số và phép tính", "Chủ đề 2: Đo lường", "Chủ đề 3: Hình học", "Chủ đề 4: Giải toán có lời văn"],
    "Lớp 3": ["Chủ đề 1: Số và phép tính", "Chủ đề 2: Đo lường", "Chủ đề 3: Hình học", "Chủ đề 4: Giải toán"],
    "Lớp 4": ["Chủ đề 1: Số tự nhiên – Phép tính", "Chủ đề 2: Phân số", "Chủ đề 3: Đo lường", "Chủ đề 4: Hình học"],
    "Lớp 5": ["Chủ đề 1: Số thập phân", "Chủ đề 2: Tỉ số – Phần trăm", "Chủ đề 3: Đo lường", "Chủ đề 4: Hình học"],
    # Lớp 6
    "Lớp 6": ["Chương 1: Số tự nhiên", "Chương 2: Số nguyên", "Chương 3: Phân số", "Chương 4: Biểu thức – Đại số", "Chương 5: Hình học trực quan"],
    # Lớp 7 (mở rộng đầy đủ chủ đề)
    "Lớp 7": [
        "Chương 1: Số hữu tỉ – Số thực",
        "Chương 2: Hàm số và đồ thị",
        "Chương 3: Hình học tam giác",
        "Chương 4: Thống kê",
        "Chương 5: Biểu thức đại số và đa thức"
    ],
    # Lớp 8 (mở rộng đầy đủ chủ đề)
    "Lớp 8": [
        "Chương 1: Đại số – Đa thức",
        "Chương 2: Phân thức",
        "Chương 3: Phương trình bậc nhất",
        "Chương 4: Hình học tứ giác – Đa giác",
        "Chương 5: Hệ phương trình và ứng dụng"
    ],
    # Lớp 9 (mở rộng đầy đủ chủ đề)
    "Lớp 9": [
        "Chương 1: Căn bậc hai – Căn thức",
        "Chương 2: Hàm số bậc nhất",
        "Chương 3: Hàm số bậc hai",
        "Chương 4: Phương trình bậc hai",
        "Chương 5: Hình học không gian – Trụ – Nón – Cầu"
    ],
}

# bai_options: giữ dạng mapping chủ đề->list bài
bai_options = {
    "Chủ đề 1: Các số đến 10": ["Đếm, đọc, viết số đến 10", "Cộng trong phạm vi 10", "Trừ trong phạm vi 10"],
    "Chủ đề 2: Các số đến 20": ["Số 11–20", "Cộng – trừ phạm vi 20"],
    "Chủ đề 3: Các số đến 100": ["Số tròn chục", "Phép tính trong phạm vi 100"],
    "Chủ đề 4: Hình học và đo lường": ["Hình tam giác – tròn – vuông – chữ nhật", "Độ dài – cm", "Thời gian – giờ"],
    "Chủ đề 5: Giải toán": ["Giải toán một bước", "Tìm số còn thiếu"],
    "Chủ đề 1: Số và phép tính": ["Số đến 100", "Cộng – trừ có nhớ", "Nhân – chia (làm quen)"],
    "Chủ đề 2: Đo lường": ["Độ dài (m, cm)", "Khối lượng (kg, g)", "Tiền Việt Nam"],
    "Chủ đề 3: Hình học": ["Góc vuông – không vuông", "Tứ giác đơn giản"],
    "Chủ đề 4: Giải toán có lời văn": ["Bài toán 1 bước", "Bài toán 2 bước"],
    "Chương 1: Số tự nhiên – Phép tính": ["Số đến 100 000", "Nhân – chia nhiều chữ số"],
    "Chương 2: Phân số": ["So sánh phân số", "Phân số bằng nhau"],
    "Chương 3: Đo lường": ["Đơn vị đo diện tích", "Diện tích hình chữ nhật – vuông"],
    "Chương 4: Hình học": ["Hình bình hành", "Hình thoi"],
    "Chương 1: Số thập phân": ["Đọc – viết số thập phân", "Tính với số thập phân"],
    "Chương 2: Tỉ số – Phần trăm": ["Tỉ số", "Tỉ lệ phần trăm"],
    "Chương 3: Đo lường (THPT)": ["Thể tích", "Diện tích hình thang – tam giác"],
    "Chương 4: Hình học (THPT)": ["Hình trụ", "Hình cầu"],
    "Chương 1: Số tự nhiên": ["Tập hợp số tự nhiên", "Chia hết – dấu hiệu chia hết"],
    "Chương 2: Số nguyên": ["Số nguyên âm – dương", "Thứ tự trong Z"],
    "Chương 3: Phân số": ["So sánh phân số", "Quy đồng phân số"],
    "Chương 4: Biểu thức – Đại số": ["Biểu thức chứa chữ", "Giá trị biểu thức"],
    "Chương 5: Hình học trực quan": ["Góc", "Tam giác"],
    # Lớp 7 mở rộng
    "Chương 1: Số hữu tỉ – Số thực": ["Số hữu tỉ", "Số thực và ký hiệu"],
    "Chương 2: Hàm số và đồ thị": ["Khái niệm hàm số", "Vẽ đồ thị hàm số đơn giản"],
    "Chương 3: Hình học tam giác": ["Tính chất tam giác", "Định lý cơ bản"],
    "Chương 4: Thống kê": ["Thu thập và biểu diễn dữ liệu", "Trung bình cộng"],
    "Chương 5: Biểu thức đại số và đa thức": ["Cộng trừ đa thức", "Nhân đa thức"],
    # Lớp 8 mở rộng
    "Chương 1: Đại số – Đa thức": ["Đa thức", "Đa thức một biến"],
    "Chương 2: Phân thức": ["Rút gọn phân thức", "Quy đồng phân thức"],
    "Chương 3: Phương trình bậc nhất": ["Phương trình cơ bản", "Ứng dụng"],
    "Chương 4: Hình học tứ giác – Đa giác": ["Tứ giác", "Đa giác"],
    "Chương 5: Hệ phương trình và ứng dụng": ["Hệ phương trình 2 ẩn", "Phương pháp thế"],
    # Lớp 9 mở rộng
    "Chương 1: Căn bậc hai – Căn thức": ["Định nghĩa căn", "Rút gọn căn"],
    "Chương 2: Hàm số bậc nhất": ["Độ dốc và giao điểm", "Ứng dụng thực tế"],
    "Chương 3: Hàm số bậc hai": ["Đỉnh và trục đối xứng", "Giải bài toán bằng đồ thị"],
    "Chương 4: Phương trình bậc hai": ["Nghiệm và hệ số", "Giải bằng công thức"],
    "Chương 5: Hình học không gian – Trụ – Nón – Cầu": ["Thể tích trụ", "Diện tích xung quanh nón"]
}

# --- Sidebar: giao diện và tham số ma trận ---
with st.sidebar:
    st.header("Thông tin sinh đề")
    lop = st.selectbox("Chọn lớp", lop_options, index=5 if len(lop_options) > 5 else 0)
    chuong_list = chuong_options.get(lop, [])
    # Nếu có chuong_list thì cho phép chọn nhiều chủ đề (multi-select)
    if chuong_list:
        chuong_selected = st.multiselect("Chọn chủ đề / chương (có thể chọn nhiều):", chuong_list, default=[chuong_list[0]])
        # flatten selected to string for prompt (join bằng '; ')
        chuong_display = "; ".join(chuong_selected) if chuong_selected else ""
    else:
        # nếu không có dữ liệu, cho phép nhập thủ công (single)
        chuong_selected = []
        chuong_display = st.text_input("Chưa có chủ đề cho lớp này, nhập thủ công:", "")

    # Build bai_list từ tất cả chủ đề được chọn
    bai_list_combined = []
    for ch in chuong_selected:
        items = bai_options.get(ch, [])
        for it in items:
            if it not in bai_list_combined:
                bai_list_combined.append(it)

    if bai_list_combined:
        bai_selected = st.multiselect("Chọn bài (có thể chọn nhiều):", bai_list_combined, default=[bai_list_combined[0]])
        bai_display = "; ".join(bai_selected) if bai_selected else ""
    else:
        bai_selected = []
        bai_display = st.text_input("Chưa có bài cho chủ đề này, nhập thủ công:", "")

    st.markdown("---")
    st.subheader("⚙️ Phân bổ theo Ma trận (CV 7991 Tối giản)")

    # Cấu hình số câu hỏi tổng cộng
    so_cau = st.number_input("Tổng số câu hỏi", min_value=1, max_value=50, value=21)

    col_nl, col_ds, col_tl = st.columns(3)
    with col_nl:
        phan_bo_nl = st.number_input("NL (Nhiều Lựa chọn)", min_value=0, value=12)
    with col_ds:
        phan_bo_ds = st.number_input("DS (Đúng - Sai)", min_value=0, value=2)
    with col_tl:
        phan_bo_tl = st.number_input("TL (Tự luận/Trả lời ngắn)", min_value=0, value=7)

    st.markdown("---")
    st.subheader("Độ khó (Cognitive Level)")

    col_nb, col_th, col_vd = st.columns(3)
    with col_nb:
        so_cau_nb = st.number_input("Nhận biết", min_value=0, value=6)
    with col_th:
        so_cau_th = st.number_input("Thông hiểu", min_value=0, value=8)
    with col_vd:
        so_cau_vd = st.number_input("Vận dụng/VDC", min_value=0, value=7)

    total_check = int(phan_bo_nl + phan_bo_ds + phan_bo_tl)
    total_level = int(so_cau_nb + so_cau_th + so_cau_vd)

    if total_check != so_cau:
        st.error(f"Tổng số câu theo loại (NL+DS+TL) = {total_check} không khớp Tổng ({so_cau}).")
    if total_level != so_cau:
        st.error(f"Tổng cấp độ (NB+TH+VĐ) = {total_level} không khớp Tổng ({so_cau}).")

    co_dap_an = st.checkbox("Có đáp án", value=True)

# --- BUILD PROMPT ---
def build_prompt(lop, chuong, bai, so_cau,
                 phan_bo_nl, phan_bo_ds, phan_bo_tl,
                 so_cau_nb, so_cau_th, so_cau_vd, co_dap_an):
    """
    Trả về prompt (string). Hàm nhận đúng 11 tham số tương ứng với generate_questions.
    Note: chuong và bai có thể là chuỗi chứa nhiều mục, đã được nối bằng '; ' ở trên.
    """
    dan_ap = "Tạo Đáp án và Lời giải chi tiết sau mỗi câu hỏi." if co_dap_an else "Không cần Đáp án."
    prompt_ma_tran = f"""
Cấu trúc ĐỀ VÀ MA TRẬN ĐẶC TẢ TỐI GIẢN (Tổng {so_cau} câu):
1. PHẦN TRẮC NGHIỆM KHÁCH QUAN (NL/DS)
    - Số câu Nhiều Lựa chọn (NL): {phan_bo_nl} câu.
    - Số câu Đúng - Sai (DS): {phan_bo_ds} câu.
2. PHẦN TỰ LUẬN (TL) / TRẢ LỜI NGẮN
    - Số câu Tự luận/Trả lời ngắn (TL): {phan_bo_tl} câu.

PHÂN BỔ MỨC ĐỘ NHẬN THỨC:
    - Nhận biết: {so_cau_nb} câu
    - Thông hiểu: {so_cau_th} câu
    - Vận dụng/VDC: {so_cau_vd} câu

YÊU CẦU ĐỀ BÀI:
1. Tạo {so_cau} câu hỏi, trong đó:
    - {phan_bo_nl} câu Trắc nghiệm 4 lựa chọn (A, B, C, D).
    - {phan_bo_ds} câu Trắc nghiệm Đúng - Sai (mỗi câu có 4 ý a, b, c, d).
    - {phan_bo_tl} câu Tự luận hoặc Trả lời ngắn.
2. Đảm bảo tổng số câu theo từng mức độ nhận thức (NB/TH/VĐ) khớp với phân bổ trên.
3. Đặt Tiêu đề rõ ràng cho từng phần.
4. Mỗi câu hỏi phải được gắn nhãn Mức độ và Loại câu hỏi (ví dụ: Câu 1. [NL - Nhận biết]).
5. Toàn bộ công thức toán phải được viết bằng LaTeX và **phải** đặt trong delimiters $$...$$. Ví dụ: $$\\frac{{a}}{{b}}$$
6. {dan_ap}
"""
    prompt_context = f"""
Bạn là giáo viên Toán, hãy sinh đề kiểm tra cho {lop} theo sách "Kết nối tri thức với cuộc sống".
- Chủ đề/Chương: {chuong}
- Bài: {bai}
{prompt_ma_tran}
"""
    return prompt_context

# --- GỌI API (Google Generative Language) ---
def generate_questions(api_key, lop, chuong, bai, so_cau,
                       phan_bo_nl, phan_bo_ds, phan_bo_tl,
                       so_cau_nb, so_cau_th, so_cau_vd, co_dap_an):
    MODEL = "models/gemini-2.5-flash"
    url = f"https://generativelanguage.googleapis.com/v1/{MODEL}:generateContent?key={api_key}"

    prompt = build_prompt(lop, chuong, bai, so_cau,
                          phan_bo_nl, phan_bo_ds, phan_bo_tl,
                          so_cau_nb, so_cau_th, so_cau_vd, co_dap_an)

    payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}

    headers = {"Content-Type": "application/json"}
    try:
        r = requests.post(url, json=payload, headers=headers, timeout=60)
        if r.status_code != 200:
            try:
                j_error = r.json()
                error_message = j_error.get("error", {}).get("message", r.text)
            except Exception:
                error_message = r.text
            return False, f"❌ Lỗi API {r.status_code}: {error_message}"
        j = r.json()
        # Cố gắng lấy text trong cấu trúc trả về
        if j.get("candidates") and len(j["candidates"]) > 0:
            cand = j["candidates"][0]
            content = cand.get("content", {})
            parts = content.get("parts", [])
            if parts and len(parts) > 0:
                text = parts[0].get("text", "")
                return True, text
        # Fallback: nếu response khác
        return False, "❌ Lỗi: AI không trả về nội dung hợp lệ."
    except requests.exceptions.Timeout:
        return False, "❌ Lỗi kết nối: Yêu cầu hết thời gian chờ (Timeout)."
    except Exception as e:
        logger.error(traceback.format_exc())
        return False, f"❌ Lỗi kết nối hoặc xử lý dữ liệu: {e}"

# --- LaTeX handling ---
LATEX_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)

def find_latex_blocks(text):
    return [(m.span(), m.group(0), m.group(1)) for m in LATEX_RE.finditer(text)]

def render_latex_png_bytes(latex_code, fontsize=20, dpi=200):
    """
    Render latex_code (no $$) thành PNG bytes. Trả về None nếu render thất bại.
    """
    try:
        # Estimate figure size based on length of latex_code (để không quá bé)
        fig = plt.figure(figsize=(1, 1))
        fig.patch.set_alpha(0.0)
        txt = f"${latex_code}$"
        t = fig.text(0.0, 0.5, txt, fontsize=fontsize, va='center', ha='left')
        fig.tight_layout(pad=0.1)
        buf = io.BytesIO()
        plt.axis('off')
        plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', pad_inches=0.05, transparent=True)
        plt.close(fig)
        buf.seek(0)
        data = buf.read()
        if not data or len(data) < 100:
            logger.warning("render_latex_png_bytes: ảnh quá nhỏ hoặc rỗng.")
            return None
        return data
    except Exception as e:
        logger.error("Lỗi render LaTeX: %s", e)
        logger.debug(traceback.format_exc())
        return None

# --- Tạo DOCX ---
def create_docx_bytes(text):
    doc = Document()
    last = 0
    latex_blocks = find_latex_blocks(text)
    for span, full, inner in latex_blocks:
        start, end = span
        before = text[last:start]
        # thêm văn bản trước công thức (giữ nguyên dòng)
        for line in before.splitlines():
            # tránh thêm dòng rỗng không cần thiết
            doc.add_paragraph(line)
        # render latex
        png_bytes = render_latex_png_bytes(inner)
        if png_bytes:
            try:
                img_stream = io.BytesIO(png_bytes)
                p = doc.add_paragraph()
                r = p.add_run()
                # Điều chỉnh width hợp lý
                r.add_picture(img_stream, width=Inches(4.0))
            except Exception as e:
                logger.error("Lỗi chèn ảnh vào DOCX: %s", e)
                doc.add_paragraph(full)
        else:
            doc.add_paragraph(full)
        last = end
    # phần còn lại
    for line in text[last:].splitlines():
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# --- Tạo PDF ---
def create_pdf_bytes(text):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    margin = 40
    y = height - 50
    line_height = 14

    def check_page():
        nonlocal y
        if y < margin + 20:
            c.showPage()
            y = height - 50

    last = 0
    latex_blocks = find_latex_blocks(text)
    for span, full, inner in latex_blocks:
        start, end = span
        before = text[last:start]
        for line in before.splitlines():
            check_page()
            safe_line = line
            max_chars = 90
            while len(safe_line) > 0:
                chunk = safe_line[:max_chars]
                c.drawString(margin, y, html.unescape(chunk))
                y -= line_height
                safe_line = safe_line[max_chars:]
                check_page()
        # render latex -> image
        png_bytes = render_latex_png_bytes(inner)
        if png_bytes:
            try:
                img_reader = ImageReader(io.BytesIO(png_bytes))
                img = Image.open(io.BytesIO(png_bytes))
                draw_w = min(400, width - 2 * margin)
                draw_h = img.height / img.width * draw_w
                check_page()
                if draw_h > (y - margin):
                    draw_h = y - margin
                    draw_w = img.width / img.height * draw_h
                c.drawImage(img_reader, margin, y - draw_h, width=draw_w, height=draw_h, mask='auto')
                y -= draw_h + 8
            except Exception as e:
                logger.error("Lỗi chèn LaTeX vào PDF: %s", e)
                check_page()
                c.drawString(margin, y, full)
                y -= line_height
        else:
            check_page()
            c.drawString(margin, y, full)
            y -= line_height
        last = end

    for line in text[last:].splitlines():
        check_page()
        safe_line = line
        max_chars = 90
        while len(safe_line) > 0:
            chunk = safe_line[:max_chars]
            c.drawString(margin, y, html.unescape(chunk))
            y -= line_height
            safe_line = safe_line[max_chars:]
            check_page()

    c.save()
    buf.seek(0)
    return buf

# --- Utility: sanitize filename ---
def sanitize_filename(s):
    # remove accents, keep ascii letters, numbers, underscore, dash
    if not s:
        return "file"
    s = str(s)
    s = unicodedata.normalize('NFKD', s)
    s = s.encode('ascii', 'ignore').decode('ascii')
    s = re.sub(r'[^\w\s-]', '', s).strip().lower()
    s = re.sub(r'[-\s]+', '_', s)
    if not s:
        return "file"
    return s[:120]

# --- BUTTON xử lý chính ---
if st.button("🎯 Sinh đề ngay", type="primary", use_container_width=True):
    # Kiểm tra điều kiện ma trận trước khi gọi API
    if not api_key:
        st.error("Thiếu API Key! Vui lòng nhập khóa API của bạn.")
    elif total_check != so_cau or total_level != so_cau:
        st.error("Lỗi Ma trận: Tổng số câu theo loại (NL/DS/TL) hoặc theo cấp độ (NB/TH/VĐ) phải bằng Tổng số câu.")
    else:
        # Chuẩn bị chuong_display và bai_display (nếu người dùng đã nhập text thủ công)
        if isinstance(chuong_display, str):
            chuong_for_prompt = chuong_display
        else:
            # fallback
            chuong_for_prompt = "; ".join(chuong_selected) if chuong_selected else ""

        if isinstance(bai_display, str):
            bai_for_prompt = bai_display
        else:
            bai_for_prompt = "; ".join(bai_selected) if bai_selected else ""

        with st.spinner("⏳ AI đang tạo đề dựa trên Ma trận Đặc tả..."):
            ok, result = generate_questions(api_key, lop, chuong_for_prompt, bai_for_prompt, so_cau,
                                           phan_bo_nl, phan_bo_ds, phan_bo_tl,
                                           so_cau_nb, so_cau_th, so_cau_vd, co_dap_an)
        if not ok:
            st.error(result)
        else:
            st.success("🎉 Đã tạo xong đề theo Ma trận Đặc tả. (Hiển thị nội dung).")
            st.markdown("---")
            st.subheader("Nội dung Đề (Raw Text)")
            # hiển thị an toàn: convert newlines thành <br> (escape trước để tránh html injection)
            safe_html = html.escape(result).replace("\n", "<br>")
            st.markdown(safe_html, unsafe_allow_html=True)
            st.markdown("---")

            latex_blocks = find_latex_blocks(result)
            download_col1, download_col2, download_col3 = st.columns(3)

            # Tên file cơ bản
            base_name = f"De_{lop}_{chuong_for_prompt}_{bai_for_prompt}"
            base_name_sanitized = sanitize_filename(base_name)

            if not latex_blocks:
                st.warning("Không tìm thấy công thức LaTeX ( $$...$$ ). Chỉ có thể xuất raw TXT.")
                with download_col1:
                    st.download_button(
                        "📥 Tải TXT", data=result.encode("utf-8"),
                        file_name=f"{base_name_sanitized}.txt", mime="text/plain",
                        use_container_width=True
                    )
            else:
                # Tạo DOCX
                try:
                    docx_io = create_docx_bytes(result)
                    with download_col1:
                        st.download_button(
                            "📥 Tải DOCX (công thức là ảnh)",
                            data=docx_io.getvalue(),
                            file_name=f"{base_name_sanitized}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                except Exception as e:
                    logger.error("Lỗi tạo DOCX: %s", traceback.format_exc())
                    with download_col1:
                        st.error(f"Lỗi tạo DOCX: {e}")

                # Tạo PDF
                try:
                    pdf_io = create_pdf_bytes(result)
                    with download_col2:
                        st.download_button(
                            "📥 Tải PDF (công thức là ảnh)",
                            data=pdf_io.getvalue(),
                            file_name=f"{base_name_sanitized}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                except Exception as e:
                    logger.error("Lỗi tạo PDF: %s", traceback.format_exc())
                    with download_col2:
                        st.error(f"Lỗi tạo PDF: {e}")

                # Luôn có TXT để kiểm tra nhanh
                with download_col3:
                    st.download_button(
                        "📥 Tải TXT (raw)",
                        data=result.encode("utf-8"),
                        file_name=f"{base_name_sanitized}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )

# --- Footer / hướng dẫn nhanh ---
st.markdown("---")
st.markdown(
    """
    **Ghi chú nhanh:**
    - Có thể chọn **nhiều chủ đề** và **nhiều bài** cùng lúc (dành cho khi bạn muốn đề bao quát nhiều nội dung).
    - Mọi công thức toán học **phải** được tạo dưới dạng `$$ ... $$` để hệ thống chuyển sang ảnh trong DOCX/PDF.
    - Nếu gặp lỗi API: kiểm tra `GOOGLE_API_KEY` hoặc message lỗi hiển thị.
    - Muốn mở rộng danh sách chương/bài: chỉnh `chuong_options` và `bai_options` trong mã.
    """
)
