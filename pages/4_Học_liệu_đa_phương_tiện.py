import requests
import streamlit as st
import time
from io import BytesIO
from docx import Document
from gtts import gTTS

# ===============================
# 1. CẤU HÌNH TRANG
# ===============================
st.set_page_config(
    page_title="Trợ lý Giáo dục AI (Gemini)",
    layout="wide",
    page_icon="🎓"
)

st.title("🎓 Trợ lý Giáo dục Đa năng (Gemini AI)")

# ===============================
# 2. NHẬP GOOGLE API KEY
# ===============================
with st.expander("🔑 Hướng dẫn lấy Google API Key"):
    st.markdown("""
    1. Truy cập: https://aistudio.google.com/app/apikey  
    2. Đăng nhập Gmail  
    3. Nhấn **Create API key** 4. Copy API Key  
    ⚠️ Không chia sẻ API Key
    """)

api_key = st.text_input("🔐 Google API Key", type="password")

if not api_key:
    st.warning("⚠️ Vui lòng nhập API Key để bắt đầu")
    st.stop()

# ===============================
# 3. DỮ LIỆU CHƯƠNG – BÀI (ĐẦY ĐỦ)
# ===============================
chuong_options_lop = {
    "Lớp 6": [
        "Chương I: Tập hợp các số tự nhiên",
        "Chương II: Tính chia hết trong tập hợp các số tự nhiên",
        "Chương III: Số nguyên",
        "Chương IV: Một số hình phẳng trong thực tiễn",
        "Chương V: Tính đối xứng của hình phẳng trong tự nhiên",
        "Chương VI: Phân số",
        "Chương VII: Số thập phân",
        "Chương VIII: Những hình hình học cơ bản",
        "Chương IX: Dữ liệu và xác suất thực nghiệm",
        "Hoạt động thực hành trải nghiệm"
    ],
    "Lớp 7": [
        "Chương I: Số hữu tỉ",
        "Chương II: Số thực",
        "Chương III: Góc và đường thẳng song song",
        "Chương IV: Tam giác bằng nhau",
        "Chương V: Thu thập và biểu diễn dữ liệu",
        "Chương VI: Tỉ lệ thức và đại lượng tỉ lệ",
        "Chương VII: Biểu thức đại số và đa thức một biến",
        "Chương VIII: Làm quen với biến cố và xác suất",
        "Chương IX: Quan hệ giữa các yếu tố trong một tam giác",
        "Chương X: Một số hình khối trong thực tiễn",
        "Bài tập ôn tập cuối năm"
    ],
    "Lớp 8": [
        "Chương I: Đa thức",
        "Chương II: Hằng đẳng thức đáng nhớ và ứng dụng",
        "Chương III: Tứ giác",
        "Chương IV: Định lí Thalès",
        "Chương V: Dữ liệu và biểu đồ",
        "Chương VI: Phân thức đại số",
        "Chương VII: Phương trình bậc nhất và hàm số bậc nhất",
        "Chương VIII: Mở đầu về tính xác suất của biến cố",
        "Chương IX: Tam giác đồng dạng",
        "Chương X: Một số hình khối trong thực tiễn",
        "Bài tập ôn tập cuối năm"
    ],
    "Lớp 9": [
        "Chương I: Phương trình và hệ hai phương trình bậc nhất hai ẩn",
        "Chương II: Phương trình và bất phương trình bậc nhất một ẩn",
        "Chương III: Căn bậc hai và căn bậc ba",
        "Chương IV: Hệ thức lượng trong tam giác vuông",
        "Chương V: Đường tròn",
        "Hoạt động thực hành trải nghiệm",
        "Chương VI: Hàm số y = ax^2 (a khác 0). Phương trình bậc hai một ẩn",
        "Chương VII: Tần số và tần số tương đối",
        "Chương VIII: Xác suất của biến cố trong một số mô hình xác suất đơn giản",
        "Chương IX: Đường tròn ngoại tiếp và đường tròn nội tiếp",
        "Chương X: Một số hình khối trong thực tiễn"
    ]
}

# --- Từng bài chi tiết ---
bai_options_lop = {
    "Lớp 6": {
        "Chương I: Tập hợp các số tự nhiên": ["Bài 1","Bài 2","Bài 3","Bài 4","Ôn tập"],
        "Chương II: Tính chia hết trong tập hợp các số tự nhiên": ["Bài 5","Bài 6","Ôn tập"],
        "Chương III: Số nguyên": ["Bài 7","Bài 8","Ôn tập"],
        "Chương IV: Một số hình phẳng trong thực tiễn": ["Bài 9","Bài 10","Ôn tập"],
        "Chương V: Tính đối xứng của hình phẳng trong tự nhiên": ["Bài 11","Bài 12","Ôn tập"],
        "Chương VI: Phân số": ["Bài 13","Bài 14","Ôn tập"],
        "Chương VII: Số thập phân": ["Bài 15","Bài 16","Ôn tập"],
        "Chương VIII: Những hình hình học cơ bản": ["Bài 17","Bài 18","Ôn tập"],
        "Chương IX: Dữ liệu và xác suất thực nghiệm": ["Bài 19","Bài 20","Ôn tập"],
        "Hoạt động thực hành trải nghiệm": ["Bài 21","Bài 22","Ôn tập"]
    },
    "Lớp 7": {
        "Chương I: Số hữu tỉ": ["Bài 1. Tập hợp các số hữu tỉ","Bài 2. Cộng, trừ, nhân, chia số hữu tỉ","Bài 3. Luỹ thừa với số mũ tự nhiên của một số hữu tỉ","Bài 4. Thứ tự thực hiện các phép tính. Quy tắc chuyển vế","Ôn tập chương I"],
        "Chương II: Số thực": ["Bài 5. Làm quen với số thập phân vô hạn tuần hoàn","Bài 6. Số vô tỉ. Căn bậc hai số học","Bài 7. Tập hợp các số thực","Ôn tập chương II"],
        "Chương III: Góc và đường thẳng song song": ["Bài 8. Góc ở vị trí đặc biệt. Tia phân giác của một góc","Bài 9. Hai đường thẳng song song và dấu hiệu nhận biết","Bài 10. Tiên đề Euclid. Tính chất của hai đường thẳng song song","Bài 11. Định lí và chứng minh định lí","Ôn tập chương III"],
        "Chương IV: Tam giác bằng nhau": ["Bài 12. Tổng các góc trong một tam giác","Bài 13. Hai tam giác bằng nhau. Trường hợp bằng nhau thứ nhất của tam giác","Bài 14. Trường hợp bằng nhau thứ hai và thứ ba của tam giác","Bài 15. Các trường hợp bằng nhau của tam giác vuông","Bài 16. Tam giác cân. Đường trung trực của đoạn thẳng","Ôn tập chương IV"],
        "Chương V: Thu thập và biểu diễn dữ liệu": ["Bài 17. Thu thập và phân loại dữ liệu","Bài 18. Biểu đồ hình quạt tròn","Bài 19. Biểu đồ đoạn thẳng","Ôn tập chương V"],
        "Chương VI: Tỉ lệ thức và đại lượng tỉ lệ": ["Bài 20. Tỉ lệ thức","Bài 21. Tính chất của dãy tỉ số bằng nhau","Bài 22. Đại lượng tỉ lệ thuận","Bài 23. Đại lượng tỉ lệ nghịch","Ôn tập chương VI"],
        "Chương VII: Biểu thức đại số và đa thức một biến": ["Bài 24. Biểu thức đại số","Bài 25. Đa thức một biến","Bài 26. Phép cộng và phép trừ đa thức một biến","Bài 27. Phép nhân đa thức một biến","Bài 28. Phép chia đa thức một biến","Ôn tập chương VII"],
        "Chương VIII: Làm quen với biến cố và xác suất": ["Bài 29. Làm quen với biến cố","Bài 30. Làm quen với xác suất của biến cố","Ôn tập chương VIII"],
        "Chương IX: Quan hệ giữa các yếu tố trong một tam giác": ["Bài 31. Quan hệ giữa góc và cạnh đối diện trong một tam giác","Bài 32. Quan hệ giữa đường vuông góc và đường xiên","Bài 33. Quan hệ giữa ba cạnh của một tam giác","Bài 34. Sự đồng quy của ba đường trung tuyến, ba đường phân giác trong một tam giác","Bài 35. Sự đồng quy của ba đường trung trực, ba đường cao trong một tam giác","Ôn tập chương IX"],
        "Chương X: Một số hình khối trong thực tiễn": ["Bài 36. Hình hộp chữ nhật và hình lập phương","Bài 37. Hình lăng trụ đứng tam giác và hình lăng trụ đứng tứ giác","Ôn tập chương X"],
        "Bài tập ôn tập cuối năm": []
    },
    "Lớp 8": {
        "Chương I: Đa thức": ["Bài 1. Đơn thức","Bài 2. Đa thức","Bài 3. Phép cộng và phép trừ đa thức","Bài 4. Phép nhân đa thức","Bài 5. Phép chia đa thức cho đơn thức","Ôn tập chương I"],
        "Chương II: Hằng đẳng thức đáng nhớ và ứng dụng": ["Bài 6. Hiệu hai bình phương. Bình phương của một tổng hay một hiệu","Bài 7. Lập phương của một tổng. Lập phương của một hiệu","Bài 8. Tổng và hiệu hai lập phương","Bài 9. Phân tích đa thức thành nhân tử","Ôn tập chương II"],
        "Chương III: Tứ giác": ["Bài 10. Tứ giác","Bài 11. Hình thang cân","Bài 12. Hình bình hành","Bài 13. Hình chữ nhật","Bài 14. Hình thoi và hình vuông","Ôn tập chương III"],
        "Chương IV: Định lí Thalès": ["Bài 15. Định lí Thalès trong tam giác","Bài 16. Đường trung bình của tam giác","Bài 17. Tính chất đường phân giác của tam giác","Ôn tập chương IV"],
        "Chương V: Dữ liệu và biểu đồ": ["Bài 18. Thu thập và phân loại dữ liệu","Bài 19. Biểu diễn dữ liệu bằng bảng, biểu đồ","Bài 20. Phân tích số liệu thống kê dựa vào biểu đồ","Ôn tập chương V"],
        "Chương VI: Phân thức đại số": ["Bài 21. Phân thức đại số","Bài 22. Tính chất cơ bản của phân thức đại số","Bài 23. Phép cộng và phép trừ phân thức đại số","Bài 24. Phép nhân và phép chia phân thức đại số","Ôn tập chương VI"],
        "Chương VII: Phương trình bậc nhất và hàm số bậc nhất": ["Bài 25. Phương trình bậc nhất một ẩn","Bài 26. Giải bài toán bằng cách lập phương trình","Bài 27. Khái niệm hàm số và đồ thị của hàm số","Bài 28. Hàm số bậc nhất và đồ thị của hàm số bậc nhất","Bài 29. Hệ số góc của đường thẳng","Ôn tập chương VII"],
        "Chương VIII: Mở đầu về tính xác suất của biến cố": ["Bài 30. Kết quả có thể và kết quả thuận lợi","Bài 31. Cách tính xác suất của biến cố bằng tỉ số","Bài 32. Mối liên hệ giữa xác suất thực nghiệm với xác suất và ứng dụng","Ôn tập chương VIII"],
        "Chương IX: Tam giác đồng dạng": ["Bài 33. Hai tam giác đồng dạng","Bài 34. Ba trường hợp đồng dạng của hai tam giác","Bài 35. Định lí Pythagore và ứng dụng","Bài 36. Các trường hợp đồng dạng của hai tam giác vuông","Bài 37. Hình đồng dạng","Ôn tập chương IX"],
        "Chương X: Một số hình khối trong thực tiễn": ["Bài 38. Hình chóp tam giác đều","Bài 39. Hình chóp tứ giác đều","Ôn tập chương X"],
        "Bài tập ôn tập cuối năm": []
    },
    "Lớp 9": {
        "Chương I: Phương trình và hệ hai phương trình bậc nhất hai ẩn": ["Bài 1. Khái niệm phương trình và hệ hai phương trình bậc nhất hai ẩn","Bài 2. Giải hệ hai phương trình bậc nhất hai ẩn","Luyện tập chung","Bài 3. Giải bài toán bằng cách lập hệ phương trình","Bài tập cuối chương I"],
        "Chương II: Phương trình và bất phương trình bậc nhất một ẩn": ["Bài 4. Phương trình quy về phương trình bậc nhất một ẩn","Bài 5. Bất đẳng thức và tính chất","Luyện tập chung","Bài 6. Bất phương trình bậc nhất một ẩn","Bài tập cuối chương II"],
        "Chương III: Căn bậc hai và căn bậc ba": ["Bài 7. Căn bậc hai và căn thức bậc hai","Bài 8. Khai căn bậc hai với phép nhân và phép chia","Luyện tập chung","Bài 9. Biến đổi đơn giản và rút gọn biểu thức chứa căn thức bậc hai","Bài 10. Căn bậc ba và căn thức bậc ba","Luyện tập chung","Bài tập cuối chương III"],
        "Chương IV: Hệ thức lượng trong tam giác vuông": ["Bài 11. Tỉ số lượng giác của góc nhọn","Bài 12. Một số hệ thức giữa cạnh, góc trong tam giác vuông và ứng dụng","Luyện tập chung","Bài tập cuối chương IV"],
        "Chương V: Đường tròn": ["Bài 13. Mở đầu về đường tròn","Bài 14. Cung và dây của một đường tròn","Bài 15. Độ dài của cung tròn. Diện tích hình quạt tròn và hình vành khuyên","Luyện tập chung","Bài 16. Vị trí tương đối của đường thẳng và đường tròn","Bài 17. Vị trí tương đối của hai đường tròn","Luyện tập chung","Bài tập cuối chương V"],
        "Hoạt động thực hành trải nghiệm": ["Pha chế dung dịch theo nồng độ yêu cầu","Tính chiều cao và xác định khoảng cách"],
        "Chương VI: Hàm số y = ax^2 (a khác 0). Phương trình bậc hai một ẩn": ["Bài 18. Hàm số y = ax2 (a ≠ 0)","Bài 19. Phương trình bậc hai một ẩn","Luyện tập chung","Bài 20. Định lí Viète và ứng dụng","Bài 21. Giải bài toán bằng cách lập phương trình","Luyện tập chung","Bài tập cuối chương VI"],
        "Chương VII: Tần số và tần số tương đối": ["Bài 22. Bảng tần số và biểu đồ tần số","Bài 23. Bảng tần số tương đối và biểu đồ tần số tương đối","Luyện tập chung","Bài 24. Bảng tần số, tần số tương đối ghép nhóm và biểu đồ","Bài tập cuối chương VII"],
        "Chương VIII: Xác suất của biến cố trong một số mô hình xác suất đơn giản": ["Bài 25. Phép thử ngẫu nhiên và không gian mẫu","Bài 26. Xác suất của biến cố liên quan tới phép thử","Luyện tập chung","Bài tập cuối chương VIII"],
        "Chương IX: Đường tròn ngoại tiếp và đường tròn nội tiếp": ["Bài 27. Góc nội tiếp","Bài 28. Đường tròn ngoại tiếp và đường tròn nội tiếp của một tam giác","Luyện tập chung","Bài 29. Tứ giác nội tiếp","Bài 30. Đa giác đều","Luyện tập chung","Bài tập cuối chương IX"],
        "Chương X: Một số hình khối trong thực tiễn": ["Bài 31. Hình trụ và hình nón","Bài 32. Hình cầu","Luyện tập chung","Bài tập cuối chương X"]
    }
}

# ===============================
# 4. HÀM GỌI GEMINI (Đã xử lý lỗi & Model chuẩn)
# ===============================
def generate_with_gemini(prompt, api_key, retry=3):
    MODEL = "gemini-2.5-flash" # Cập nhật lên bản 2.0 mới nhất hoặc dùng 1.5-flash
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL}:generateContent?key={api_key}"

    payload = {
        "contents": [{
            "role": "user",
            "parts": [{"text": prompt}]
        }]
    }
    
    headers = {'Content-Type': 'application/json'}

    for attempt in range(retry):
        try:
            response = requests.post(url, json=payload, headers=headers, timeout=60)
            
            if response.status_code == 200:
                data = response.json()
                if "candidates" in data and len(data["candidates"]) > 0:
                    return data["candidates"][0]["content"]["parts"][0]["text"]
                else:
                    return "⚠️ API trả về nhưng không có nội dung (Safety blocking?)."
            
            elif response.status_code == 503:
                time.sleep(2)  # Chờ rồi thử lại
            
            else:
                return f"❌ Lỗi API {response.status_code}: {response.text}"
                
        except Exception as e:
            return f"❌ Lỗi kết nối: {e}"

    return "⚠️ Gemini đang quá tải. Thầy/cô vui lòng thử lại sau 1–2 phút."

# ===============================
# 5. TẠO FILE WORD
# ===============================
def create_docx_bytes(text):
    doc = Document()
    doc.add_heading("TÀI LIỆU TOÁN HỌC (AI)", 0)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ===============================
# 6. GIAO DIỆN TABS
# ===============================
tab1, tab2, tab3 = st.tabs([
    "📘 Tổng hợp kiến thức", 
    "🎵 Nhạc Toán", 
    "🎧 Đọc văn bản"
])

# -------- TAB 1: TỔNG HỢP KIẾN THỨC (Cập nhật dịch H'Mông) ----------
with tab1:
    c1, c2, c3 = st.columns(3)
    with c1:
        lop = st.selectbox("Lớp", list(chuong_options_lop.keys()))
    with c2:
        chuong = st.selectbox("Chương", chuong_options_lop[lop])
    with c3:
        # Lấy danh sách bài, xử lý trường hợp không có dữ liệu
        ds_bai = bai_options_lop.get(lop, {}).get(chuong, ["Bài mở đầu"])
        if len(ds_bai) == 0: ds_bai = ["Ôn tập chương"] # Fallback
        bai = st.selectbox("Bài", ds_bai)

    # Nút tạo nội dung Tiếng Việt
    if st.button("🚀 Tổng hợp nội dung (Tiếng Việt)"):
        prompt = f"""
        Bạn là giáo viên Toán THCS.
        Hãy soạn nội dung cho: {bai} – thuộc {chuong} ({lop})

        Yêu cầu:
        1. Trình bày Ngắn gọn, dễ hiểu.
        2. Có: Khái niệm, Công thức, Ví dụ minh họa, 3 Bài tập tự luyện (có đáp án).
        3. Sử dụng Markdown để trình bày đẹp.
        """
        with st.spinner("⏳ Đang tạo nội dung..."):
            text = generate_with_gemini(prompt, api_key)
            st.session_state["math_content"] = text
            # Xóa bản dịch cũ nếu tạo bài mới
            if "hmong_content" in st.session_state:
                del st.session_state["hmong_content"]

    # Hiển thị nội dung Tiếng Việt nếu đã có
    if "math_content" in st.session_state:
        st.subheader("🇻🇳 Nội dung Tiếng Việt")
        st.markdown(st.session_state["math_content"])
        
        st.download_button(
            "📥 Tải file Word (Tiếng Việt)",
            create_docx_bytes(st.session_state["math_content"]),
            file_name="Toan_AI_Vietnamese.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.markdown("---")
        
        # Phần Dịch sang H'Mông
        st.subheader("🌏 Hỗ trợ ngôn ngữ vùng cao")
        col_trans_1, col_trans_2 = st.columns([1, 3])
        
        with col_trans_1:
            if st.button("🔄 Dịch sang tiếng H'Mông"):
                trans_prompt = f"""
                Bạn là một chuyên gia ngôn ngữ và giáo dục vùng cao.
                Hãy dịch toàn bộ nội dung toán học dưới đây sang tiếng H'Mông (Hmoob).
                
                Yêu cầu quan trọng:
                1. Giữ nguyên toàn bộ các công thức toán học, số liệu và định dạng Markdown/LaTeX.
                2. Dịch thuật ngữ toán học chính xác nhưng dễ hiểu cho học sinh dân tộc.
                3. Giữ nguyên cấu trúc bài (Khái niệm, Ví dụ, Bài tập).
                
                Nội dung cần dịch:
                {st.session_state["math_content"]}
                """
                with st.spinner("⏳ Đang dịch sang tiếng H'Mông..."):
                    hmong_text = generate_with_gemini(trans_prompt, api_key)
                    st.session_state["hmong_content"] = hmong_text

        # Hiển thị kết quả dịch
        if "hmong_content" in st.session_state:
            st.markdown("### 🟢 Nội dung tiếng H'Mông (Hmoob)")
            st.markdown(st.session_state["hmong_content"])
            
            st.download_button(
                "📥 Tải file Word (Tiếng H'Mông)",
                create_docx_bytes(st.session_state["hmong_content"]),
                file_name="Toan_AI_Hmong.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# -------- TAB 2: NHẠC TOÁN ----------
with tab2:
    st.header("Sáng tác Nhạc Toán học")
    style = st.selectbox("Phong cách", ["Rap", "Vè dân gian", "Pop", "Thơ lục bát"])
    if st.button("🎤 Sáng tác bài hát"):
        # Lấy thông tin từ Tab 1
        content_context = st.session_state.get("math_content", f"Bài {bai} - {chuong}")
        
        prompt = f"""
        Hãy viết lời bài hát theo phong cách {style}.
        Chủ đề: Giúp học sinh nhớ kiến thức toán học của bài: {bai} - {chuong}.
        Dựa trên nội dung chính: {content_context[:500]}...
        Yêu cầu: Vui nhộn, dễ thuộc.
        """
        with st.spinner("Đang sáng tác..."):
            st.markdown(generate_with_gemini(prompt, api_key))

# -------- TAB 3: ĐỌC VĂN BẢN (TTS) ----------
with tab3:
    st.header("Chuyển văn bản thành giọng nói")
    tts_text = st.text_area("Nhập văn bản", "Chào các em học sinh thân mến!")
    if st.button("▶️ Đọc văn bản"):
        if tts_text.strip():
            try:
                with st.spinner("Đang xử lý âm thanh..."):
                    tts = gTTS(text=tts_text, lang="vi")
                    # Dùng bộ nhớ đệm (RAM) để không phải lưu file
                    audio_bytes = BytesIO()
                    tts.write_to_fp(audio_bytes)
                    st.audio(audio_bytes, format="audio/mp3")
            except Exception as e:
                st.error(f"Lỗi: {e}")
        else:
            st.warning("Vui lòng nhập nội dung.")
